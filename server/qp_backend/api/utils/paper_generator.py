import os
import json
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from datetime import datetime
import logging
from django.conf import settings

# Configure logging
logging.basicConfig(level=logging.INFO)

class QuestionPaperGenerator:
    @staticmethod
    def format_table_cell(cell, text, alignment=WD_ALIGN_PARAGRAPH.LEFT):
        paragraph = cell.paragraphs[0]
        paragraph.alignment = alignment
        run = paragraph.add_run(text)
        return run

    @staticmethod
    def add_equation_to_docx(paragraph, equation_data):
        try:
            math_element = parse_xml(equation_data['mathml'])
            paragraph._p.append(math_element)
            run = paragraph.add_run()
            run.add_text(" ")
        except Exception as e:
            logging.error(f"Error adding equation: {e}")
            paragraph.add_run(equation_data.get('text', 'Equation Error'))

    @staticmethod
    def add_image_to_docx(doc, cell, image_path):
        try:
            paragraph = cell.paragraphs[0]
            if os.path.exists(image_path):
                run = paragraph.add_run()
                picture = run.add_picture(image_path)
                max_width = Inches(6)
                if picture.width > max_width:
                    aspect_ratio = picture.height / picture.width
                    picture.width = max_width
                    picture.height = int(max_width * aspect_ratio)
        except Exception as e:
            logging.error(f"Error adding image: {e}")
            paragraph.add_run(f"[Image: {os.path.basename(image_path)}]")

    @staticmethod
    def create_paper(metadata, selected_questions, questions_data):
        doc = Document()
        
        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Header
        header = doc.add_heading('DEPARTMENT OF INFORMATION SCIENCE AND ENGINEERING', 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Create header table
        table = doc.add_table(rows=5, cols=4)
        table.style = 'Table Grid'

        # Fill header information
        header_data = [
            ['Date', metadata.date.strftime('%d-%m-%Y'), 'Maximum Marks', str(metadata.max_marks)],
            ['Course Code', metadata.course_code, 'Duration', metadata.duration],
            ['Sem', metadata.semester, 'Improvement CIE', 'Yes' if metadata.is_improvement_cie else 'No'],
            ['UG/PG', 'UG', 'Faculty:', metadata.faculty.name],
            ['Course Title', metadata.course_title, '', '']
        ]

        for i, row_data in enumerate(header_data):
            for j, text in enumerate(row_data):
                cell = table.cell(i, j)
                QuestionPaperGenerator.format_table_cell(cell, text, WD_ALIGN_PARAGRAPH.CENTER)

        doc.add_paragraph()

        def add_questions_to_table(questions, table, start_num=1):
            total_marks = 0
            for i, selection in enumerate(questions, start_num):
                question_data = next(q for q in questions_data if q.q_id == selection.question.q_id)
                row_cells = table.add_row().cells
                
                # Question number
                QuestionPaperGenerator.format_table_cell(row_cells[0], str(i), WD_ALIGN_PARAGRAPH.CENTER)
                
                # Question content cell
                question_cell = row_cells[1]
                if not question_cell.text.strip():
                    question_cell._element.clear_content()
                
                # Add question text
                para = question_cell.add_paragraph()
                para.add_run(question_data.text)
                
                # Add equations if present
                media = question_data.media.first()
                if media and media.equations:
                    equations = media.equations
                    for equation in equations:
                        eq_para = question_cell.add_paragraph()
                        QuestionPaperGenerator.add_equation_to_docx(eq_para, equation)
                
                # Add images if present
                if media and media.image_paths:
                    for img_path in media.image_paths:
                        QuestionPaperGenerator.add_image_to_docx(doc, question_cell, img_path)
                
                # Add other fields
                QuestionPaperGenerator.format_table_cell(row_cells[2], str(question_data.marks), WD_ALIGN_PARAGRAPH.CENTER)
                QuestionPaperGenerator.format_table_cell(row_cells[3], question_data.bt, WD_ALIGN_PARAGRAPH.CENTER)
                QuestionPaperGenerator.format_table_cell(row_cells[4], question_data.co, WD_ALIGN_PARAGRAPH.CENTER)
                
                total_marks += question_data.marks
            
            return total_marks

        # Part A
        doc.add_heading('Part- A', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
        partA_table = doc.add_table(rows=1, cols=5)
        partA_table.style = 'Table Grid'
        
        # Set header row
        header_cells = partA_table.rows[0].cells
        headers = ['Q. No.', 'Questions', 'M', 'BT', 'CO']
        for i, text in enumerate(headers):
            QuestionPaperGenerator.format_table_cell(header_cells[i], text, WD_ALIGN_PARAGRAPH.CENTER)

        # Add Part A questions
        partA_questions = [q for q in selected_questions if q.part == 'A']
        partA_marks = add_questions_to_table(partA_questions, partA_table)

        doc.add_paragraph()

        # Part B
        doc.add_heading('Part- B', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
        partB_table = doc.add_table(rows=1, cols=5)
        partB_table.style = 'Table Grid'
        
        header_cells = partB_table.rows[0].cells
        for i, text in enumerate(headers):
            QuestionPaperGenerator.format_table_cell(header_cells[i], text, WD_ALIGN_PARAGRAPH.CENTER)

        # Add Part B questions
        partB_questions = [q for q in selected_questions if q.part == 'B']
        partB_marks = add_questions_to_table(partB_questions, partB_table)

        # Add footer with total marks calculation
        doc.add_paragraph()
        footer_para = doc.add_paragraph("*********")
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("BT-Blooms Taxonomy, CO-Course Outcomes")
        
        total_marks_para = doc.add_paragraph(f"Total Marks: {partA_marks + partB_marks}")
        total_marks_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        return doc 